"""skills_extra: encryption, archive, conversion, and webhook skills.

Extracted from services/skills_extra.py.
"""
from __future__ import annotations
import csv
import io
import json
import os
import re
import shutil
import time
import zipfile
from datetime import datetime
from pathlib import Path
from typing import Any

from app.core.config import DATA_DIR, GENERATED_DIR, UPLOAD_DIR

OUTPUT_DIR = GENERATED_DIR
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
WORKSPACE = DATA_DIR / "workspace"
WORKSPACE.mkdir(parents=True, exist_ok=True)
BACKEND_UPLOADS = UPLOAD_DIR


# в•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђ
# 1. РЁРР¤Р РћР’РђРќРР• (Fernet = AES-128-CBC)
# в•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђ

_KEY_FILE = DATA_DIR / "elira_secret.key"


def _get_fernet():
    try:
        from cryptography.fernet import Fernet
    except ImportError:
        return None, "pip install cryptography"
    if _KEY_FILE.exists():
        key = _KEY_FILE.read_bytes()
    else:
        key = Fernet.generate_key()
        _KEY_FILE.write_bytes(key)
    return Fernet(key), None


def encrypt_text(text: str) -> dict:
    f, err = _get_fernet()
    if err:
        return {"ok": False, "error": err}
    token = f.encrypt(text.encode("utf-8"))
    return {"ok": True, "encrypted": token.decode("utf-8"), "original_length": len(text)}


def decrypt_text(token: str) -> dict:
    f, err = _get_fernet()
    if err:
        return {"ok": False, "error": err}
    try:
        plain = f.decrypt(token.encode("utf-8"))
        return {"ok": True, "decrypted": plain.decode("utf-8")}
    except Exception as e:
        return {"ok": False, "error": f"Р Р°СЃС€РёС„СЂРѕРІРєР° РЅРµ СѓРґР°Р»Р°СЃСЊ: {e}"}


# в•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђ
# 2. РђР РҐРР’РђРўРћР 
# в•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђ

def create_zip(source_path: str, output_name: str = "") -> dict:
    """Р—Р°РїР°РєРѕРІС‹РІР°РµС‚ С„Р°Р№Р»/РїР°РїРєСѓ РІ ZIP."""
    src = Path(source_path)
    if not src.exists():
        # РџРѕРїСЂРѕР±СѓРµРј РІ workspace
        src = WORKSPACE / source_path
    if not src.exists():
        return {"ok": False, "error": f"РќРµ РЅР°Р№РґРµРЅРѕ: {source_path}"}

    fname = output_name or f"{src.stem}_{int(time.time())}.zip"
    if not fname.endswith(".zip"):
        fname += ".zip"
    out = OUTPUT_DIR / fname

    with zipfile.ZipFile(out, "w", zipfile.ZIP_DEFLATED) as zf:
        if src.is_file():
            zf.write(src, src.name)
        else:
            for f in src.rglob("*"):
                if f.is_file() and ".git" not in str(f) and "__pycache__" not in str(f):
                    zf.write(f, f.relative_to(src))

    return {"ok": True, "path": str(out), "filename": fname, "size": out.stat().st_size,
            "download_url": f"/api/skills/download/{fname}"}


def extract_zip(zip_path: str, dest: str = "") -> dict:
    """Р Р°СЃРїР°РєРѕРІС‹РІР°РµС‚ ZIP."""
    zp = Path(zip_path)
    if not zp.exists():
        return {"ok": False, "error": f"РќРµ РЅР°Р№РґРµРЅ: {zip_path}"}
    dest_dir = Path(dest) if dest else WORKSPACE / zp.stem
    dest_dir.mkdir(parents=True, exist_ok=True)

    with zipfile.ZipFile(zp, "r") as zf:
        zf.extractall(dest_dir)
        names = zf.namelist()

    return {"ok": True, "dest": str(dest_dir), "files": names[:50], "count": len(names)}


# в•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђ
# 3. РљРћРќР’Р•Р РўР•Р  Р¤РђР™Р›РћР’
# в•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђ

def convert_file(source_path: str, target_format: str) -> dict:
    """РљРѕРЅРІРµСЂС‚РёСЂСѓРµС‚: CSVв†’XLSX, JSONв†’CSV, MDв†’DOCX, XLSXв†’CSV."""
    src = Path(source_path)
    if not src.exists():
        src = WORKSPACE / source_path
        if not src.exists():
            src = BACKEND_UPLOADS / source_path
    if not src.exists():
        return {"ok": False, "error": f"РќРµ РЅР°Р№РґРµРЅ: {source_path}"}

    ext = src.suffix.lower()
    target = target_format.lower().strip(".")
    fname = f"{src.stem}.{target}"
    out = OUTPUT_DIR / fname

    try:
        if ext == ".csv" and target == "xlsx":
            return _csv_to_xlsx(src, out, fname)
        elif ext == ".json" and target == "csv":
            return _json_to_csv(src, out, fname)
        elif ext == ".md" and target == "docx":
            return _md_to_docx(src, out, fname)
        elif ext in (".xlsx", ".xls") and target == "csv":
            return _xlsx_to_csv(src, out, fname)
        else:
            return {"ok": False, "error": f"РљРѕРЅРІРµСЂС‚Р°С†РёСЏ {ext} в†’ .{target} РЅРµ РїРѕРґРґРµСЂР¶РёРІР°РµС‚СЃСЏ"}
    except Exception as e:
        return {"ok": False, "error": str(e)}


def _csv_to_xlsx(src, out, fname):
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill
    wb = Workbook()
    ws = wb.active
    with open(src, encoding="utf-8", errors="replace") as f:
        reader = csv.reader(f)
        for r, row in enumerate(reader, 1):
            for c, val in enumerate(row, 1):
                cell = ws.cell(row=r, column=c, value=val)
                if r == 1:
                    cell.font = Font(bold=True)
                    cell.fill = PatternFill(start_color="D5E8F0", end_color="D5E8F0", fill_type="solid")
    wb.save(str(out))
    return {"ok": True, "filename": fname, "download_url": f"/api/skills/download/{fname}"}


def _json_to_csv(src, out, fname):
    data = json.loads(src.read_text(encoding="utf-8"))
    if isinstance(data, dict):
        data = [data]
    if not isinstance(data, list) or not data:
        return {"ok": False, "error": "JSON РґРѕР»Р¶РµРЅ Р±С‹С‚СЊ РјР°СЃСЃРёРІРѕРј РѕР±СЉРµРєС‚РѕРІ"}
    keys = list(data[0].keys())
    with open(out, "w", encoding="utf-8", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=keys)
        writer.writeheader()
        for row in data:
            writer.writerow(row)
    return {"ok": True, "filename": fname, "download_url": f"/api/skills/download/{fname}"}


def _md_to_docx(src, out, fname):
    from docx import Document
    doc = Document()
    text = src.read_text(encoding="utf-8")
    for line in text.split("\n"):
        s = line.strip()
        if s.startswith("# "):
            doc.add_heading(s[2:], level=1)
        elif s.startswith("## "):
            doc.add_heading(s[3:], level=2)
        elif s.startswith("### "):
            doc.add_heading(s[4:], level=3)
        elif s.startswith("- ") or s.startswith("* "):
            doc.add_paragraph(s[2:], style="List Bullet")
        elif s:
            doc.add_paragraph(s)
    doc.save(str(out))
    return {"ok": True, "filename": fname, "download_url": f"/api/skills/download/{fname}"}


def _xlsx_to_csv(src, out, fname):
    from openpyxl import load_workbook
    wb = load_workbook(src, read_only=True, data_only=True)
    ws = wb.active
    with open(out, "w", encoding="utf-8", newline="") as f:
        writer = csv.writer(f)
        for row in ws.iter_rows(values_only=True):
            writer.writerow(row)
    wb.close()
    return {"ok": True, "filename": fname, "download_url": f"/api/skills/download/{fname}"}


# в•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђ
# 4. REGEX РџРћРњРћР©РќРРљ
# в•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђ

def test_regex(pattern: str, text: str, flags: str = "") -> dict:
    """РўРµСЃС‚РёСЂСѓРµС‚ regex РЅР° СЃС‚СЂРѕРєРµ."""
    try:
        fl = 0
        if "i" in flags:
            fl |= re.IGNORECASE
        if "m" in flags:
            fl |= re.MULTILINE
        if "s" in flags:
            fl |= re.DOTALL

        matches = []
        for m in re.finditer(pattern, text, fl):
            matches.append({
                "match": m.group(),
                "start": m.start(),
                "end": m.end(),
                "groups": list(m.groups()) if m.groups() else [],
            })

        return {
            "ok": True,
            "pattern": pattern,
            "text": text,
            "matches": matches,
            "count": len(matches),
            "has_match": len(matches) > 0,
        }
    except re.error as e:
        return {"ok": False, "error": f"РќРµРІР°Р»РёРґРЅС‹Р№ regex: {e}"}


# в•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђ
# 5. РџР•Р Р•Р’РћР”Р§РРљ (С‡РµСЂРµР· Ollama)
# в•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђ

def translate_text(text: str, target_lang: str = "english", model: str = "qwen3:8b") -> dict:
    """РџРµСЂРµРІРѕРґ С‡РµСЂРµР· LLM."""
    try:
        import ollama
        resp = ollama.chat(
            model=model,
            messages=[{
                "role": "user",
                "content": f"Translate the following text to {target_lang}. Output ONLY the translation, nothing else.\n\n{text}"
            }],
            options={"temperature": 0.3, "num_predict": 2048},
        )
        translated = resp.get("message", {}).get("content", "").strip()
        return {"ok": True, "original": text, "translated": translated, "target_lang": target_lang, "model": model}
    except Exception as e:
        return {"ok": False, "error": str(e)}


# в•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђ
# 6. CSV / Р”РђРќРќР«Р• РђРќРђР›РР—
# в•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђ

def analyze_csv(file_path: str, query: str = "") -> dict:
    """РђРЅР°Р»РёР·РёСЂСѓРµС‚ CSV С„Р°Р№Р»: СЃС‚Р°С‚РёСЃС‚РёРєР°, РїРµСЂРІС‹Рµ СЃС‚СЂРѕРєРё, Р°РіСЂРµРіР°С†РёРё."""
    fp = Path(file_path)
    if not fp.exists():
        fp = WORKSPACE / file_path
        if not fp.exists():
            fp = BACKEND_UPLOADS / file_path
    if not fp.exists():
        return {"ok": False, "error": f"РќРµ РЅР°Р№РґРµРЅ: {file_path}"}

    try:
        import pandas as pd
        df = pd.read_csv(fp, encoding="utf-8", on_bad_lines="skip")

        result = {
            "ok": True,
            "filename": fp.name,
            "shape": {"rows": len(df), "columns": len(df.columns)},
            "columns": list(df.columns),
            "dtypes": {col: str(dtype) for col, dtype in df.dtypes.items()},
            "head": df.head(5).to_dict(orient="records"),
            "describe": {},
            "nulls": df.isnull().sum().to_dict(),
        }

        # РЎС‚Р°С‚РёСЃС‚РёРєР° РїРѕ С‡РёСЃР»РѕРІС‹Рј РєРѕР»РѕРЅРєР°Рј
        num_cols = df.select_dtypes(include=["int64", "float64"]).columns
        if len(num_cols) > 0:
            desc = df[num_cols].describe()
            result["describe"] = desc.to_dict()

        # Р•СЃР»Рё РµСЃС‚СЊ Р·Р°РїСЂРѕСЃ вЂ” РІС‹РїРѕР»РЅСЏРµРј eval
        if query.strip():
            try:
                eval_result = df.eval(query) if not query.strip().startswith("df") else eval(query, {"df": df, "pd": pd})
                if hasattr(eval_result, "to_dict"):
                    result["query_result"] = eval_result.head(20).to_dict(orient="records") if hasattr(eval_result, "head") else eval_result.to_dict()
                else:
                    result["query_result"] = str(eval_result)
            except Exception as e:
                result["query_error"] = str(e)

        return result
    except ImportError:
        return {"ok": False, "error": "pip install pandas"}
    except Exception as e:
        return {"ok": False, "error": str(e)}


# в•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђ
# 7. WEBHOOK РҐР РђРќРР›РР©Р•
# в•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђ

_webhooks: list[dict] = []
_MAX_WEBHOOKS = 100


def store_webhook(data: dict, source: str = "unknown") -> dict:
    global _webhooks
    entry = {
        "id": len(_webhooks),
        "received_at": datetime.utcnow().isoformat() + "Z",
        "source": source,
        "data": data,
    }
    _webhooks.append(entry)
    if len(_webhooks) > _MAX_WEBHOOKS:
        _webhooks = _webhooks[-_MAX_WEBHOOKS:]
    return {"ok": True, "id": entry["id"]}


def list_webhooks(limit: int = 20) -> dict:
    return {"ok": True, "items": _webhooks[-limit:], "count": len(_webhooks)}


def clear_webhooks() -> dict:
    global _webhooks
    _webhooks = []
    return {"ok": True}

